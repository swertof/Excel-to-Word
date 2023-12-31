from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
import time
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Menu

# Глобальная переменная для хранения имени файла Excel
excel_file_name = ""
def browse_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if file_path:
        workbook = openpyxl.load_workbook(file_path)
        global sheet
        sheet = workbook.active
        excel_file_name = file_path  # Обновляем имя файла
        browse_button.config(text=f"Обзор ({excel_file_name})")
def show_error_message():
    messagebox.showerror("Ошибка", "Укажите путь к Excel файлу")
    global error
    error=True
doc = Document("Образец.docx")


#Функции изменения полей в документе
def input_values(sample,new,underline):
    if new == None:
        new = ""
    count = 0
    for paragraph in doc.paragraphs:
        if count > 0:
            break
        if sample in paragraph.text:
            count += 1
            for run in paragraph.runs:
                if sample in run.text:
                    # Заменяем образец текста новым текстом с сохранением форматирования
                    run.text = run.text.replace(sample, str(new), 1)
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    run.underline = underline
                    
def input_values_into_table(sample, new):
    if new == None:
        new = ""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if count > 0:
                    break
                if sample in cell.text:
                    count += 1
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if sample in run.text:
                                # Заменяем образец текста новым текстом с сохранением форматирования
                                run.text = run.text.replace(sample, str(new),1)
                                run.font.size = Pt(10)
                                run.font.name = "Times New Roman"

def inf_input_values_into_table(sample, new):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if sample in cell.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if sample in run.text:
                                # Заменяем образец текста новым текстом с сохранением форматирования
                                run.text = run.text.replace(sample, new)
                                run.font.size = Pt(10)
                                run.font.name = "Times New Roman"

def input_LS(new):
    #Ввод лицевого счета
    sample = "/ЛС"
    underline = True
    input_values(sample, new, underline)
def input_sity(new):
    #Ввод города
    sample = "/Город"
    underline = True
    input_values(sample, new, underline)
def input_street(new):
    sample = "/Улица"
    underline = True
    input_values(sample, new, underline)
def input_house(new):
    #Ввод номера дома
    sample = "/НомерДома"
    underline = True
    input_values(sample, new, underline)
def input_apartments(new):
    #Ввод номера квартиры
    sample = "/НомерКв"
    underline = True
    input_values(sample, new, underline)
def input_consumerName(new):
    #Ввод ФИО потребителя
    sample = "{X X X X}"
    underline = True
    input_values(sample, new, underline)
def input_removedType(new):
    #Ввод типа снятого ПУ
    sample = "/ТипСнят"
    input_values_into_table(sample, new)
def input_removedNumber(new):
    #Ввод номера снятого ПУ
    sample = "/НомерСнят"
    input_values_into_table(sample, new)
def input_installedType(new):
    #Ввод типа нового ПУ
    sample = "/ТипУст"
    input_values_into_table(sample, new)
def input_installedNumber(new):
    #Ввод номера нового ПУ
    sample = "/НомерУст"
    input_values_into_table(sample, new)
def input_plomb(new):
    #Ввод установленной пломбы
    sample = "/Пломба"
    input_values_into_table(sample, new)
def input_authorName(new):
    #Ввод ФИО
    sample = "/ФИО"
    underline = True
    input_values(sample, new, underline)

def plomb_counting(plomb,i):
    try:
        plomb=plomb.split("*")
        num_part = plomb[1][:2]
        num=plomb[1][2::]
        number = plomb[0] + "*" + num_part + str(int(num) + i)
        return number
    except:
        pass

LS_col = "A"
street_col = "B"
house_col = "C"
apartment_col = "D"
consumerName_col = "E"
removeNum_col = "H"
removeType_col = "I"
def process_data():
    global special_installed
    try:
        plomb = plomb_entry.get()
        city = city_entry.get()
        common_installedType = installedType_entry.get()
        special_installed = special_installed_entry.get().split()
        authorName = authorName_entry.get()

        if special_installed:
            special_installedtype = special_installedtype_entry.get()
        
        street = sheet[f'{street_col}{3}'].value
        house = sheet[f'{house_col}{3}'].value
        # Счетчик строк
        # Первая строка не пустая!
        global non_empty_row_count
        non_empty_row_count = 0
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                non_empty_row_count += 1
        j = 0
        for i in range(2, non_empty_row_count+1):
            ls = sheet[f'{LS_col}{i}'].value
            apartments = sheet[f'{apartment_col}{i}'].value
            consumer = sheet[f'{consumerName_col}{i}'].value
            removeNum = sheet[f'{removeNum_col}{i}'].value
            removeType = sheet[f'{removeType_col}{i}'].value

            installedType = common_installedType
            if str(apartments) in special_installed:
                installedType = special_installedtype

            input_consumerName(consumer)
            input_LS(ls)
            input_sity(city)
            input_street(street)
            input_house(house)
            input_apartments(apartments)
            input_removedType(removeType)
            input_removedNumber(removeNum)
            input_installedType(installedType)
            #input_plomb(plomb_counting(plomb,j))
            j += 1
            input_authorName(authorName)
            global error
            error=False
    except NameError:
        show_error_message()
root = tk.Tk()
root.title("Замена текста в документе Word")
root.geometry("400x400")

# Создаем и размещаем виджеты на первой странице
frame1 = ttk.Frame(root)


browse_button = ttk.Button(frame1, text="Обзор", command=browse_file)
city_label = ttk.Label(frame1, text="Город:")
city_entry = ttk.Entry(frame1)
installedType_label = ttk.Label(frame1, text="Тип установленных ПУ:")
installedType_entry = ttk.Entry(frame1)
special_installed_label = ttk.Label(frame1, text="Номера GPRS ПУ через пробел\n(Например: 1 11 21)):")
special_installed_entry = ttk.Entry(frame1)
special_installedtype_label = ttk.Label(frame1, text="Тип GPRS ПУ:")
special_installedtype_entry = ttk.Entry(frame1)
authorName_label = ttk.Label(frame1, text="ФИО представителя:")
authorName_entry = ttk.Entry(frame1)
plomb_label = ttk.Label(frame1, text="Пломба")
plomb_entry = ttk.Entry(frame1)
page=1
def nextPage_and_process():
        process_data()
        if error == False:
            frame1.pack_forget()
            frame2.pack()
            root.after(500, show_first)
            global page
            page += 1
        
def enter_for_buttons():
    if page == 1:
        nextPage_and_process()
    elif page == 2:
        range_handler()
next_button = ttk.Button(frame1, text="Далее", command=nextPage_and_process)
def wanna_change_cols():
    frame1.pack_forget()
    frame3.pack(expand=True, fill='both')
cols_change = ttk.Button(frame1, text="Изменить номера столбцов", command=wanna_change_cols)
root.bind("<Return>", lambda event=None: enter_for_buttons())

browse_button.pack()
city_label.pack()
city_entry.pack()
plomb_label.pack()
plomb_entry.pack()
installedType_label.pack()
installedType_entry.pack()
special_installed_label.pack()
special_installed_entry.pack()
special_installedtype_label.pack()
special_installedtype_entry.pack()
authorName_label.pack()
authorName_entry.pack()
next_button.pack()
cols_change.pack()


i = 2
def show_first():
    apartments = sheet[f'{apartment_col}{i}'].value
    removeNum = sheet[f'{removeNum_col}{i}'].value
    apartments_label.config(text=f"Квартира:{apartments}")
    removeNum_label.config(text=f"Снят ПУ:{removeNum}")

def input_export():
    # Заблокировать кнопку ввода и поле ввода
    installedNum_entry.config(state=tk.DISABLED)
    enter_button.config(state=tk.DISABLED)
    global i
    apartments_label.config(text=f"Квартира:{sheet[f'{apartment_col}{i+1}'].value}")
    removeNum_label.config(text=f"Снят ПУ:{sheet[f'{removeNum_col}{i+1}'].value}")
    installedNum = installedNum_entry.get()
    input_installedNumber(installedNum)
    i += 1
    
    #Разблокировать кнопку после выполнения
    installedNum_entry.config(state=tk.NORMAL)
    enter_button.config(state=tk.NORMAL)
def save():
    #При сохранении незаполненного файла все образцы номеров стираются
    inf_InstaledNum=inf_InstaledNum_entry.get()
    inf_InstaledNum_special=inf_InstaledNum_special_entry.get()
    for i in range(2, non_empty_row_count+1):
            apartments=sheet[f'{apartment_col}{i}'].value
            if str(apartments) in special_installed:
                input_values_into_table("/НомерУст", inf_InstaledNum_special)
            else:
                input_values_into_table("/НомерУст", inf_InstaledNum)
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    if file_path:
        doc.save(file_path)
    root.destroy()
# Создаем фрейм
frame2 = ttk.Frame(root)
apartments_label = ttk.Label(frame2, text="")
removeNum_label = ttk.Label(frame2, text="")



def range_handler():
    if i < non_empty_row_count:
        input_export()
        installedNum_entry.delete(0, tk.END)
    else:
        input_export()
        save()

# Кнопка для ввода данных
enter_button = ttk.Button(frame2, text="Ввод", command=range_handler)
save_button = ttk.Button(frame2, text="Сохранить", command=save)
# Метка и поле ввода для номера установленного ПУ
installedNum_label = ttk.Label(frame2, text="Номер установленного ПУ:")
installedNum_entry = ttk.Entry(frame2)
installedNum_value = tk.StringVar()
inf_InstaledNum_label = ttk.Label(frame2, text="Вставить значение во все незаполненные номера установленных ПУ")
inf_InstaledNum_common_label = ttk.Label(frame2, text="Для обычных ПУ")
inf_InstaledNum_entry=ttk.Entry(frame2)
inf_InstaledNum_special_label = ttk.Label(frame2, text="Для GPRS ПУ")
inf_InstaledNum_special_entry=ttk.Entry(frame2)

# Размещаем виджеты на фрейме
apartments_label.pack()
removeNum_label.pack()
installedNum_label.pack()
installedNum_entry.pack()
enter_button.pack()
inf_InstaledNum_label.pack()
inf_InstaledNum_common_label.pack()
inf_InstaledNum_entry.pack()
inf_InstaledNum_special_label.pack()
inf_InstaledNum_special_entry.pack()
save_button.pack()

frame3 = ttk.Frame(root)

LS_col_label=ttk.Label(frame3, text="Лицевой счет")
street_col_label=ttk.Label(frame3, text="Улица")
house_col_label=ttk.Label(frame3, text="Дом")
apartment_col_label=ttk.Label(frame3, text="Номер квартиры")
consumerName_col_label=ttk.Label(frame3, text="ФИО потребителя")
removeNum_col_label=ttk.Label(frame3, text="Номер снятого ПУ")
removeType_col_label=ttk.Label(frame3, text="Тип снятого ПУ")

LS_col_entry=ttk.Entry(frame3)
street_col_entry=ttk.Entry(frame3)
house_col_entry=ttk.Entry(frame3)
apartment_col_entry=ttk.Entry(frame3)
consumerName_col_entry=ttk.Entry(frame3)
removeNum_col_entry=ttk.Entry(frame3)
removeType_col_entry=ttk.Entry(frame3)

def change():
    global LS_col, street_col, house_col, apartment_col, consumerName_col, removeNum_col, removeType_col
    LS_col=LS_col_entry.get()
    street_col=street_col_entry.get()
    house_col=house_col_entry.get()
    apartment_col=apartment_col_entry.get()
    consumerName_col=consumerName_col_entry.get()
    removeNum_col=removeNum_col_entry.get()
    removeType_col=removeType_col_entry.get()
    frame3.pack_forget()
    frame1.pack()
def cancel():
    frame3.pack_forget()
    frame1.pack()
change_button = ttk.Button(frame3, text="Сохранить", command=change)
cancel_button = ttk.Button(frame3, text="Отмена", command=cancel)

LS_col_label.pack()
LS_col_entry.pack()
street_col_label.pack()
street_col_entry.pack()
house_col_label.pack()
house_col_entry.pack()
apartment_col_label.pack()
apartment_col_entry.pack()
consumerName_col_label.pack()
consumerName_col_entry.pack()
removeNum_col_label.pack()
removeNum_col_entry.pack()
removeType_col_label.pack()
removeType_col_entry.pack()
change_button.pack()
cancel_button.pack()


# Размещаем фрейм на главном окне

root.bind("<Return>", lambda event=None: enter_for_buttons())
frame1.pack()
root.mainloop()
